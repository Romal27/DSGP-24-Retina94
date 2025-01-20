import requests
from bs4 import BeautifulSoup
from docx import Document

# URL of the website to scrape
url = "https://www.doc.lk/search?doctor=&hospital=0&specialization=14&date=&search=1"

# Set up headers for the request
headers = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
}

# Make a GET request to fetch the HTML content
response = requests.get(url, headers=headers)

# Check if the request was successful
if response.status_code == 200:
    soup = BeautifulSoup(response.content, 'html.parser')

    # Create a Word document
    document = Document()

    # Example: Extracting data (adjust selectors as needed)
    hospital_name = "Asiri Hospital - Kandy"
    document.add_heading(hospital_name, level=1)

    # Find all doctor details (adjust based on the site's HTML structure)
    doctors = soup.find_all('div', class_='doctor-details')  # Adjust class name

    for doctor in doctors:
        name = doctor.find('h3', class_='doctor-name').text.strip()  # Adjust selector
        specialization = doctor.find('p', class_='specialization').text.strip()  # Adjust selector

        # Add details to the document
        document.add_paragraph(name)
        document.add_paragraph(specialization)

    # Save the document
    output_file = "Doctors_List.docx"
    document.save(output_file)
    print(f"Data saved to {output_file}")
else:
    print(f"Failed to fetch the page. Status code: {response.status_code}")
